import streamlit as st
import pandas as pd
import anthropic
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Pillar II GloBE", page_icon="P2", layout="wide")

FREE_FIELD_COUNT = 20

JUR_INFO = {
    "ZA": {"name": "South Africa", "role": "HQ / UPE", "std_rate": 27.0},
    "AU": {"name": "Australia", "role": "Subsidiary", "std_rate": 30.0},
    "IE": {"name": "Ireland", "role": "Subsidiary", "std_rate": 12.5},
}

def init_state():
    free_zeros = {f"Free{i:02d}": 0.0 for i in range(1, FREE_FIELD_COUNT + 1)}
    if "entities" not in st.session_state:
        st.session_state.entities = pd.DataFrame([
            {**{"Entity": "OUTsurance Holdings Ltd", "Jurisdiction": "ZA", "Type": "Insurance", "Revenue": 8500.0, "PBT": 2100.0, "CoveredTaxes": 567.0, "DeferredTaxAdj": 45.0, "Payroll": 320.0, "TangibleAssets": 1200.0, "Active": True}, **free_zeros},
            {**{"Entity": "OUTsurance Life Ltd", "Jurisdiction": "ZA", "Type": "Life Insurance", "Revenue": 3200.0, "PBT": 890.0, "CoveredTaxes": 240.0, "DeferredTaxAdj": 18.0, "Payroll": 85.0, "TangibleAssets": 340.0, "Active": True}, **free_zeros},
            {**{"Entity": "OUTsurance Australia Pty", "Jurisdiction": "AU", "Type": "Gen. Insurance", "Revenue": 1800.0, "PBT": 420.0, "CoveredTaxes": 126.0, "DeferredTaxAdj": 12.0, "Payroll": 55.0, "TangibleAssets": 180.0, "Active": True}, **free_zeros},
            {**{"Entity": "OUTsurance Ireland Ltd", "Jurisdiction": "IE", "Type": "Reinsurance", "Revenue": 950.0, "PBT": 280.0, "CoveredTaxes": 35.0, "DeferredTaxAdj": 8.0, "Payroll": 22.0, "TangibleAssets": 95.0, "Active": True}, **free_zeros},
        ])
    else:
        for i in range(1, FREE_FIELD_COUNT + 1):
            col = f"Free{i:02d}"
            if col not in st.session_state.entities.columns:
                st.session_state.entities[col] = 0.0
    if "transactions" not in st.session_state:
        st.session_state.transactions = pd.DataFrame([
            {"Description": "Reinsurance premium ZA to IE", "From": "OUTsurance Holdings Ltd", "To": "OUTsurance Ireland Ltd", "Amount": 120.0, "Type": "Reinsurance Premium", "TP_Method": "TNMM", "ArmsLength": True},
            {"Description": "Management fee ZA to AU", "From": "OUTsurance Holdings Ltd", "To": "OUTsurance Australia Pty", "Amount": 45.0, "Type": "Management Fee", "TP_Method": "CUP", "ArmsLength": True},
            {"Description": "IT licence fee ZA to IE", "From": "OUTsurance Holdings Ltd", "To": "OUTsurance Ireland Ltd", "Amount": 30.0, "Type": "Royalty / IP", "TP_Method": "TNMM", "ArmsLength": False},
        ])
    if "api_key" not in st.session_state:
        st.session_state.api_key = ""
    if "free_field_labels" not in st.session_state:
        st.session_state.free_field_labels = {f"Free{i:02d}": f"Custom Field {i}" for i in range(1, FREE_FIELD_COUNT + 1)}
    if "free_field_visible" not in st.session_state:
        st.session_state.free_field_visible = {f"Free{i:02d}": False for i in range(1, FREE_FIELD_COUNT + 1)}
    if "core_field_labels" not in st.session_state:
        st.session_state.core_field_labels = {
            "Entity": "Entity Name",
            "Jurisdiction": "Jurisdiction",
            "Type": "Entity Type",
            "Revenue": "Revenue (ZARm)",
            "PBT": "Profit Before Tax (ZARm)",
            "CoveredTaxes": "Covered Taxes (ZARm)",
            "DeferredTaxAdj": "Deferred Tax Adj (ZARm)",
            "Payroll": "Payroll (ZARm)",
            "TangibleAssets": "Tangible Assets (ZARm)",
        }

init_state()

def calc_globe():
    active = st.session_state.entities[st.session_state.entities["Active"] == True].copy()
    results = []
    for jur, grp in active.groupby("Jurisdiction"):
        globe_income = (grp["PBT"] - grp["DeferredTaxAdj"]).sum()
        adj_cov_tax = (grp["CoveredTaxes"] + grp["DeferredTaxAdj"]).sum()
        sbie = (0.05 * grp["Payroll"] + 0.05 * grp["TangibleAssets"]).sum()
        net_income = max(0, globe_income - sbie)
        etr = (adj_cov_tax / globe_income * 100) if globe_income > 0 else 0
        top_up_rate = max(0, 15 - etr)
        top_up_tax = (top_up_rate / 100) * net_income
        j = JUR_INFO.get(jur, {"name": jur, "role": "", "std_rate": 0})
        status = "Compliant" if etr >= 15 else ("Marginal" if top_up_rate < 1 else "EXPOSURE")
        results.append({
            "Jur": jur, "Name": j["name"], "Role": j["role"],
            "Entities": len(grp), "GloBE_Income": globe_income,
            "Adj_Cov_Tax": adj_cov_tax, "SBIE": sbie,
            "ETR": etr, "TopUp_Rate": top_up_rate,
            "TopUp_Tax": top_up_tax,
            "QDMTT": "Yes" if top_up_rate > 0 else "No",
            "Status": status,
        })
    return pd.DataFrame(results)

def call_claude(prompt):
    key = st.session_state.get("api_key", "")
    if not key:
        try:
            key = st.secrets["ANTHROPIC_API_KEY"]
        except Exception:
            return "No API key set. Add your Anthropic API key in the sidebar."
    try:
        client = anthropic.Anthropic(api_key=key)
        msg = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}]
        )
        return msg.content[0].text
    except Exception as e:
        return "API Error: " + str(e)

def get_summary():
    results = calc_globe()
    lines = []
    for _, r in results.iterrows():
        lines.append(r["Name"] + " (" + r["Jur"] + "): GloBE Income ZAR" + str(round(r["GloBE_Income"], 1)) + "m, ETR " + str(round(r["ETR"], 2)) + "%, Top-up Tax ZAR" + str(round(r["TopUp_Tax"], 1)) + "m")
    return "\n".join(lines)

def generate_word_doc(title, content):
    doc = Document()
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("OUTsurance Group | Pillar II GloBE Compliance | FY 2024")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], 4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], 3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], 2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], 1)
        elif stripped in ("", "---"):
            doc.add_paragraph()
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def build_template_prompt(template_name, entity_list, summary, tx_list):
    base = (
        f"You are a senior international tax expert specialising in OECD Pillar II GloBE rules, "
        f"South African tax law (SARS), and insurance sector taxation.\n\n"
        f"Generate a professional, detailed {template_name} for OUTsurance Group — a South African "
        f"non-life insurance holding company (UPE) with subsidiaries in Australia and Ireland.\n\n"
        f"LIVE GROUP DATA:\nEntities: {entity_list}\n\nGloBE Calculation Results:\n{summary}\n\n"
        f"Intercompany Transactions: {tx_list}\n\n"
        "SOUTH AFRICAN (SARS) REGULATORY FRAMEWORK:\n"
        "- Income Tax Act No. 58 of 1962 (ITA): GloBE IIR enacted via Section 9H; QDMTT via Section 9I (effective 1 January 2024)\n"
        "- Taxation Laws Amendment Act (TLAB) 2023: primary GloBE enabling legislation\n"
        "- Tax Administration Act No. 28 of 2011 (TAA): Section 257B — CbCR obligations for UPEs with group revenue > ZAR 10 billion\n"
        "- Transfer Pricing: Section 31 ITA and SARS Practice Note 7 of 2005 (arm's length standard, contemporaneous documentation)\n"
        "- SARS Interpretation Note 80: deferred tax accounting under IAS 12\n"
        "- SARS Binding General Rulings BGR 14 and BGR 31: short-term insurance specific tax treatment\n"
        "- SARS IT14 corporate income tax return and IT14SD Supplementary Declaration reconciliation\n"
        "- SARS eFiling XML submission requirements and GIR notification obligations\n"
        "- Financial Sector Conduct Authority (FSCA) and Prudential Authority (South African Reserve Bank) regulatory context\n"
        "- King IV Report on Corporate Governance for South Africa 2016 — Principle 15 (responsible tax)\n"
        "- SAICA guidance notes on Pillar II GloBE implementation for South African entities\n\n"
        "GLOBAL BEST PRACTICE FRAMEWORK:\n"
        "- OECD GloBE Model Rules (December 2021) and Commentary (March 2022)\n"
        "- OECD Administrative Guidance: February 2023 (deferred tax, QDMTT), July 2023 (STTR, insurance), December 2023 (GIR filing)\n"
        "- OECD GloBE Information Return (GIR) XML schema and filing instructions (December 2023)\n"
        "- OECD Transitional CbCR Safe Harbour guidance (December 2022): De Minimis, Simplified ETR, Routine Profits tests\n"
        "- OECD BEPS Action 13: CbCR, Master File, Local File requirements\n"
        "- IFRS 17 Insurance Contracts: impact on GloBE financial accounts basis\n"
        "- IFRS 9 Financial Instruments and IAS 12 Income Taxes: deferred tax adjustments under GloBE\n"
        "- Australia Treasury Laws Amendment Act 2024: Pillar Two IIR and QDMTT implementation\n"
        "- Ireland Finance Act 2024: Section 111AAK QDMTT and IIR (15% top-up on 12.5% STR — primary group exposure)\n"
        "- Big Four technical guidance: KPMG Pillar Two Navigator, Deloitte Global Minimum Tax Analyser, PwC Pillar Two Hub, EY Global Tax Alert\n"
        "- Global insurance peer ETR benchmarks: Zurich Insurance (~15.2%), Munich Re (~16.8%), Allianz (~18.1%), Sanlam (~18.9%)\n\n"
    )
    specifics = {
        "GloBE Information Return (GIR)": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (GIR):\n"
            "- Structure per OECD GIR December 2023 XML schema: Part I (Overview), Part II (Constituent Entity Data), "
            "Part III (Jurisdiction Data), Part IV (Top-up Tax Allocation), Part V (Additional Information)\n"
            "- All mandatory data fields per OECD GIR Article 44 filing obligation\n"
            "- South Africa as UPE: first-year transitional filing deadline 18 months post fiscal year-end; "
            "subsequent years 15 months\n"
            "- SARS Section 257B TAA: notify SARS of UPE filing status and confirm no local surrogate filing\n"
            "- Constituent entity table: all 4 OUTsurance entities with TIN, jurisdiction, role, entity type\n"
            "- Jurisdiction-level summary tables: GloBE Income, Adjusted Covered Taxes, ETR, SBIE, Top-up Tax\n"
            "- QDMTT credit mechanism: Ireland Finance Act 2024 QDMTT creditable against IIR charged by ZA UPE\n"
            "- Transitional safe harbour elections: document per-jurisdiction decisions in Part V\n"
            "- Insurance-specific disclosures: IFRS 17 technical provisions treatment, DAC elimination impact\n"
            "- SARS eFiling XML submission preparation steps and IT14 cross-reference\n"
        ),
        "QDMTT Calculation Worksheet": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (QDMTT Worksheet):\n"
            "- Step-by-step calculation per OECD Article 10.1 definition of Qualified Domestic Minimum Top-up Tax\n"
            "- Ireland QDMTT (primary exposure): Finance Act 2024 Section 111AAK — calculate exact top-up from 12.5% STR to 15% GloBE minimum\n"
            "- South Africa QDMTT: Section 9I ITA — calculate to confirm nil exposure (27% STR well above 15%)\n"
            "- Ireland step-by-step: GloBE Income → Deferred Tax Adj → Adj Covered Taxes → SBIE → Net GloBE Income → ETR → Top-up Rate → QDMTT Liability\n"
            "- QDMTT safe harbour eligibility: OECD Administrative Guidance July 2023 — assess whether Ireland QDMTT qualifies for credit\n"
            "- Qualified QDMTT criteria: must meet OECD consistency, standstill and equivalence conditions\n"
            "- Deferred tax adjustments: SARS Interpretation Note 80 vs IFRS IAS 12 basis — reconciliation table\n"
            "- Insurance-specific: equalisation reserve deductibility under SARS BGR 14 vs IFRS 17 (no equalisation reserve)\n"
            "- Payment mechanics: Ireland CT1 return integration and instalment timing\n"
            "- SARS IT14SD integration: QDMTT amount in supplementary schedule, ZAR conversion at spot rate\n"
        ),
        "IIR Allocation Memorandum": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (IIR Memorandum):\n"
            "- South Africa UPE charging mechanism: Section 9H ITA and OECD Article 2.1 Income Inclusion Rule\n"
            "- Top-down allocation: OUTsurance Holdings Ltd (ZA UPE) charges IIR on low-taxed constituent entities\n"
            "- Ireland primary analysis: 12.5% STR → Finance Act 2024 QDMTT credits against IIR → residual IIR if QDMTT not fully qualifying\n"
            "- Ownership structure: direct ownership ZA → IE (100%) and ZA → AU (100%) — no intermediate parent\n"
            "- OECD Article 2.4: no intermediate parent entity complication — UPE charges IIR directly\n"
            "- OECD Article 2.6: UTPR backstop — South Africa's UTPR status as UPE jurisdiction (Section 9J ITA)\n"
            "- Section 9D CFC provisions: interaction between South Africa CFC rules and GloBE — anti-duplication\n"
            "- QDMTT credit: Ireland QDMTT reduces IIR charged by ZA on Ireland income per OECD Article 2.1(d)\n"
            "- Tax accounting entries: IAS 12 current and deferred tax for IIR in OUTsurance Holdings IT14\n"
            "- Filing obligations: SARS IT14 IIR schedule — 12 months post year-end filing\n"
            "- Elections and notifications: UPE filing status, safe harbour elections, QDMTT credit claims\n"
        ),
        "Transfer Pricing Impact Analysis": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (TP Analysis):\n"
            "- Section 31 ITA arm's length standard and SARS Practice Note 7/2005 — full compliance review\n"
            "- OECD Transfer Pricing Guidelines 2022 (July 2022 consolidated version) — Chapter I-III method selection\n"
            "- TRANSACTION 1: Reinsurance premium ZA→IE ZAR 120m (TNMM) — actuarial pricing benchmarks, OECD Chapter VI special measures for risk\n"
            "- TRANSACTION 2: Management fee ZA→AU ZAR 45m (CUP) — comparable uncontrolled service charge analysis\n"
            "- TRANSACTION 3: IT licence royalty ZA→IE ZAR 30m (TNMM) — FLAGGED NON-ARMS-LENGTH — URGENT REMEDIATION REQUIRED\n"
            "- GloBE interaction: TP adjustments under Section 31 ITA affect GloBE income per OECD Article 3.2 (financial accounts basis)\n"
            "- SARS documentation: Master File and Local File per Practice Note 7 Section 6; contemporaneous requirement\n"
            "- SARS CbCR: Section 257B TAA — group revenue > ZAR 10bn — mandatory annual CbCR filing\n"
            "- Penalty exposure: Section 75B ITA — 200% understatement penalty; SARS TP audit priority for cross-border royalties\n"
            "- Recommended actions: urgently review IT licence royalty — consider SARS Advance Pricing Agreement (APA) under Section 76P ITA\n"
            "- Benchmarking sources: Bureau van Dijk Orbis, TP Catalyst, RoyaltyStat — interquartile range analysis\n"
        ),
        "Insurance-Specific Deferred Tax Adjustment": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (Deferred Tax Adjustment):\n"
            "- OECD GloBE Article 4.4: deferred tax adjustments methodology — inclusion of deferred tax in covered taxes\n"
            "- SARS Interpretation Note 80: IAS 12 deferred tax — SARS tax basis vs IFRS accounting basis reconciliation\n"
            "- IFRS 17 Insurance Contracts (effective 1 January 2023): major impact on deferred tax positions:\n"
            "  * Contractual Service Margin (CSM): timing difference — IFRS 17 profit release not taxable until claims paid\n"
            "  * Risk Adjustment (RA): SARS Section 28(4) vs IFRS 17 fulfilment cash flow basis\n"
            "  * Loss Component: deferred tax asset recognition — IAS 12 recoverability assessment\n"
            "- IFRS 9 Financial Instruments: investment portfolio unrealised gains — SARS eighth schedule CGT timing\n"
            "- Insurance-specific temporary differences requiring GloBE adjustment:\n"
            "  * DAC (Deferred Acquisition Costs): IFRS 17 eliminates DAC — SARS Section 11(a) still allows deduction — significant deferred tax liability\n"
            "  * IBNR/IBNER reserves: SARS Section 28(4) undiscounted vs IFRS 17 discounted fulfilment cash flows\n"
            "  * Equalisation reserves: SARS Section 28(3) deductible vs IFRS 17 not recognised — deferred tax liability\n"
            "  * Unearned premium reserve: SARS vs IFRS 17 measurement differences\n"
            "- GloBE deferred tax cap: 15% rate cap on deferred tax assets (OECD Article 4.4.1)\n"
            "- Recapture rule: deferred tax assets not utilised within 5 years — recapture into covered taxes\n"
            "- Entity-level calculation tables with SARS tax basis, IFRS carrying amount, temporary difference, DTA/DTL\n"
            "- SARS BGR 14: short-term insurance specific provisions to apply\n"
        ),
        "SBIE Workpaper": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (SBIE Workpaper):\n"
            "- OECD Article 5.3: Substance-Based Income Exclusion mechanics and policy rationale\n"
            "- CRITICAL — Apply exact OECD transitional SBIE rates:\n"
            "  * FY2024: Payroll × 9.8% + Tangible Assets × 7.8%\n"
            "  * FY2025: Payroll × 7.8% + Tangible Assets × 6.6%\n"
            "  * FY2026: Payroll × 6.4% + Tangible Assets × 5.5%\n"
            "  * FY2027: Payroll × 5.8% + Tangible Assets × 5.0%\n"
            "  * Steady-state (FY2033+): Payroll × 5.0% + Tangible Assets × 5.0%\n"
            "- Eligible payroll costs per OECD Article 5.3.2: salaries, wages, employer social security contributions — "
            "SARS SDL and UIF contributions included; director fees excluded unless employment-related\n"
            "- Eligible tangible assets per OECD Article 5.3.3: PPE net book value per IFRS (NOT SARS WDV); "
            "investment properties (IAS 40) — assess insurance float eligibility\n"
            "- Exclusions from tangible assets: financial assets, intangibles, IP, real estate held for resale\n"
            "- Insurance-specific: assess whether insurance investment properties qualify as eligible tangible assets\n"
            "- Entity-level SBIE calculation table: all 4 entities with payroll, tangible assets, SBIE amount\n"
            "- Jurisdiction blending: SBIE is calculated at jurisdictional level — aggregate ZA entities\n"
            "- SBIE optimisation: confirm election made in GIR; SBIE reduces GloBE net income (tax base)\n"
            "- South Africa: large SBIE due to HQ payroll (ZAR 320m+ZAR 85m) — quantify ZA SBIE benefit\n"
            "- Forward-looking: project SBIE under declining transitional rates to FY2033 steady-state\n"
        ),
        "Transitional Safe Harbour Assessment": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (Safe Harbour Assessment):\n"
            "- OECD December 2022 Transitional CbCR Safe Harbour — three alternative qualifying tests:\n"
            "  TEST 1 — De Minimis Test:\n"
            "  * Condition: Revenue < EUR 10m AND Profit Before Tax < EUR 1m in CbCR data\n"
            "  * Assess all three jurisdictions with EUR conversion at applicable FX rate\n"
            "  TEST 2 — Simplified ETR Test:\n"
            "  * Simplified ETR = Income Tax Expense per CbCR ÷ PBT per CbCR\n"
            "  * Thresholds: FY2024: 15%, FY2025: 16%, FY2026: 17%\n"
            "  * Expected: South Africa 27% STR — PASS; Australia 30% STR — PASS; Ireland 12.5% — LIKELY FAIL\n"
            "  TEST 3 — Routine Profits Test:\n"
            "  * Condition: PBT per CbCR ≤ SBIE calculated using CbCR payroll and tangible asset data\n"
            "- CbCR data requirements: SARS Section 257B TAA — qualified CbCR per OECD criteria\n"
            "- Qualifying CbCR definition: must meet OECD consistency, reliability and completeness criteria\n"
            "- CbCR vs GloBE basis differences: document and assess materiality for safe harbour eligibility\n"
            "- Transition period scope: fiscal years beginning before 1 July 2028\n"
            "- Jurisdiction-by-jurisdiction assessment table with test results and safe harbour conclusion\n"
            "- Election notification: safe harbour elections must be made in GIR Part V filed with SARS\n"
            "- Recommended strategy: claim safe harbour for ZA and AU; perform full GloBE calculation for IE\n"
            "- Risk: if CbCR data is unreliable, safe harbour may not be available — assess fallback position\n"
        ),
        "CbCR Reconciliation": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (CbCR Reconciliation):\n"
            "- SARS Section 257B TAA: OUTsurance qualifies as UPE — group consolidated revenue > ZAR 10bn\n"
            "- OECD BEPS Action 13 CbCR template: Table 1 (jurisdiction overview), Table 2 (entity list), Table 3 (additional info)\n"
            "- SARS eFiling: CbCR XML per OECD schema v2.0 — filing deadline 12 months after financial year-end\n"
            "- Subsidiary notifications: OUTsurance Australia (ATO) and OUTsurance Ireland (Revenue Commissioners) "
            "must notify local tax authorities of UPE CbCR filing — penalty risk if omitted\n"
            "- Key reconciliation items — CbCR data vs GloBE financial accounts basis:\n"
            "  * Revenue: gross vs net of ceded reinsurance premiums (different treatment)\n"
            "  * Profit: PBT per CbCR vs GloBE Income (deferred tax adj, IFRS 17 uplifts, excluded dividends)\n"
            "  * Tax: income tax expense per CbCR vs Adjusted Covered Taxes per GloBE\n"
            "  * Employees: CbCR headcount vs GloBE eligible payroll (FTE-weighted basis)\n"
            "  * Assets: CbCR total assets vs GloBE eligible tangible assets (PPE only, IFRS basis)\n"
            "- Transitional safe harbour suitability: assess CbCR as qualified CbCR for safe harbour tests\n"
            "- SARS audit risk: material inconsistency between CbCR and IT14 — reconcile and document\n"
            "- Three-year rolling analysis: FY2022, FY2023, FY2024 CbCR trend analysis\n"
            "- GRI 207 public disclosure: consider voluntary publication of CbCR data in integrated annual report\n"
        ),
        "Board Tax Governance Report": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (Board Tax Governance Report):\n"
            "- King IV Report on Corporate Governance for South Africa 2016:\n"
            "  * Principle 15: Governing body ensures responsible corporate citizenship including tax\n"
            "  * Recommended Practice 15.2: Tax governance policy, tax risk appetite, oversight by audit committee\n"
            "  * Audit and Risk Committee Charter: specific tax risk oversight responsibilities\n"
            "- GRI 207: Tax Standard 2019 — four disclosures: approach to tax, tax governance, stakeholder engagement, country-by-country reporting\n"
            "- SARS Voluntary Disclosure Programme (VDP): note availability for any historic non-compliance\n"
            "- OECD Tax Transparency: BEPS Action 13 public CbCR reporting — consider voluntary disclosure\n"
            "- Board-level tax risk appetite statement: define tolerance for aggressive positions on GloBE scale\n"
            "- Tax Risk Management Framework (TRMF):\n"
            "  * Risk identification matrix: Pillar II, transfer pricing, QDMTT, CbCR, FSCA compliance\n"
            "  * Risk controls: Big Four advisor reviews (EY/KPMG/Deloitte/PwC), external auditor tax sign-off\n"
            "  * Escalation procedures: material tax risks to Audit Committee and Board\n"
            "- Pillar II GloBE material disclosures for Board:\n"
            "  * Ireland QDMTT exposure: quantified ZAR amount — Finance Act 2024 compliance timeline\n"
            "  * IIR filing obligation as UPE — Section 9H ITA — South Africa first-ever GloBE filing\n"
            "  * Transfer pricing risk: IT licence royalty flagged — potential Section 31 ITA adjustment\n"
            "  * GIR first filing: 18-month transitional deadline — readiness assessment\n"
            "- Insurance sector specific: FSCA Senior Manager regime — CFO and Head of Tax accountability\n"
            "- Prudential Authority (SARB): capital adequacy impact of Pillar II tax liabilities\n"
            "- Forward-looking: FY2025/2026 tax risk horizon — SBIE declining rates, UTPR implementation watch\n"
        ),
        "BEPS Action 13 Master File Summary": (
            "TEMPLATE-SPECIFIC REQUIREMENTS (BEPS Action 13 Master File):\n"
            "- OECD BEPS Action 13 — Chapter V Transfer Pricing Documentation: Master File structure\n"
            "- SARS Practice Note 7/2005 Section 6: master file requirements for South African resident entities\n"
            "- SARS TAA Section 29: record-keeping obligations — 5-year minimum retention\n"
            "- Section 31 ITA: contemporaneous documentation mandatory at time of filing IT14\n"
            "MASTER FILE STRUCTURE:\n"
            "PART I — ORGANISATIONAL STRUCTURE:\n"
            "  * OUTsurance Group legal entity chart: OUTsurance Holdings Ltd (ZA/UPE) → OUTsurance Life Ltd (ZA) + OUTsurance Australia Pty (AU) + OUTsurance Ireland Ltd (IE)\n"
            "  * Ownership percentages, incorporation dates, tax registration numbers, tax residency\n"
            "PART II — DESCRIPTION OF MNE GROUP BUSINESS:\n"
            "  * Non-life insurance (ZA), life insurance (ZA), general insurance (AU), reinsurance (IE)\n"
            "  * Group value chain: underwriting, claims, reinsurance, investment, shared services\n"
            "  * Key functions, assets, risks (FAR analysis) per entity — economic substance assessment\n"
            "PART III — INTANGIBLES:\n"
            "  * Brand, IT systems, underwriting methodology — owned and developed in South Africa\n"
            "  * DEMPE analysis: Development, Enhancement, Maintenance, Protection, Exploitation\n"
            "  * IT licence royalty ZA→IE: arm's length review required — URGENT\n"
            "PART IV — INTERCOMPANY FINANCIAL ACTIVITIES:\n"
            "  * Reinsurance premium ZAR 120m (ZA→IE), management fee ZAR 45m (ZA→AU), IT royalty ZAR 30m (ZA→IE)\n"
            "  * TP method selection rationale per transaction\n"
            "PART V — FINANCIAL AND TAX POSITIONS:\n"
            "  * Consolidated financials reference, Pillar II GloBE ETR by jurisdiction, SARS CbCR confirmation\n"
            "- Penalty risk: SARS Section 75B ITA — 200% understatement penalty for TP non-compliance\n"
        ),
    }
    specific = specifics.get(
        template_name,
        "Format as a complete professional compliance document with numbered sections, data tables, methodology notes, rule references, and action items."
    )
    footer = (
        "\n\nDOCUMENT FORMAT REQUIREMENTS:\n"
        "1. Formal header: document title, date (March 2026), prepared by (Tax Department), reviewed by (Head of Tax), classification: CONFIDENTIAL\n"
        "2. Executive Summary (max 3 paragraphs)\n"
        "3. Numbered sections with clear headings\n"
        "4. Data tables for all quantitative calculations (columns: Description | Amount ZARm | Notes)\n"
        "5. Cite specific OECD Article numbers, SARS section references, and IFRS standards throughout\n"
        "6. Risk matrix where applicable (Risk | Likelihood | Impact | Mitigation)\n"
        "7. Action items table: Action | Owner | Deadline | Status\n"
        "8. Produce a complete, ready-to-use professional compliance document — do not truncate.\n"
    )
    return base + specific + footer


def build_benchmarking_prompt(query, summary):
    return (
        "You are a Pillar II GloBE expert and South African tax specialist advising the OUTsurance Group tax team.\n"
        "OUTsurance is a South African non-life insurance holding company (UPE) with subsidiaries in Australia and Ireland.\n\n"
        "SOUTH AFRICAN REGULATORY CONTEXT:\n"
        "- Section 9H ITA (GloBE IIR), Section 9I ITA (QDMTT), Section 9J ITA (UTPR) — effective 1 January 2024\n"
        "- TLAB 2023: primary enabling legislation; SARS Practice Note 7/2005: transfer pricing\n"
        "- Section 257B TAA: CbCR obligations; SARS Interpretation Note 80: deferred tax\n"
        "- SARS BGR 14 and BGR 31: insurance-specific tax treatment\n"
        "- King IV Principle 15: tax governance; FSCA and Prudential Authority oversight\n\n"
        "GLOBAL FRAMEWORK:\n"
        "- OECD GloBE Model Rules 2021, Commentary 2022, Administrative Guidance Feb/Jul/Dec 2023\n"
        "- OECD Transitional CbCR Safe Harbour (December 2022)\n"
        "- IFRS 17, IFRS 9, IAS 12 — insurance and tax accounting interactions\n"
        "- Australia Treasury Laws Amendment Act 2024; Ireland Finance Act 2024 (Section 111AAK QDMTT)\n"
        "- Big Four guidance: KPMG, Deloitte, PwC, EY Pillar Two trackers\n\n"
        "GROUP POSITION:\n" + summary + "\n\n"
        "QUESTION: " + query + "\n\n"
        "Provide a technical, specific answer covering:\n"
        "1. Direct answer with conclusion\n"
        "2. Relevant OECD GloBE Model Rule articles and SARS legislation section references\n"
        "3. Insurance-sector specific considerations (IFRS 17, DAC, technical provisions, reinsurance)\n"
        "4. Application to OUTsurance's three-jurisdiction structure (ZA/AU/IE)\n"
        "5. Comparison against Big Four (KPMG, Deloitte, PwC, EY) published guidance\n"
        "6. Practical action steps with responsible owners\n"
        "7. Filing deadlines and penalty risks\n\n"
        "Be specific, technical, and cite legislation and OECD rule references throughout."
    )


with st.sidebar:
    st.markdown("## OUTsurance\n### Pillar II GloBE")
    st.markdown("---")
    st.markdown("#### Anthropic API Key")
    api_input = st.text_input("API Key", value=st.session_state.api_key, type="password", placeholder="sk-ant-...", label_visibility="collapsed")
    if api_input and api_input != st.session_state.api_key:
        st.session_state.api_key = api_input
        st.success("Key saved")
    st.markdown("---")
    page = st.radio("Go to", ["Dashboard", "Entities", "Transactions", "Upload Data", "AI Templates", "Benchmarking", "Settings"], label_visibility="collapsed")
    st.markdown("---")
    r_side = calc_globe()
    ttu = r_side["TopUp_Tax"].sum() if not r_side.empty else 0
    st.metric("Est. Top-up Tax", "ZAR " + str(round(ttu, 1)) + "m")
    st.metric("Active Entities", len(st.session_state.entities[st.session_state.entities["Active"] == True]))

st.markdown("# OUTsurance - Pillar II GloBE Reporting")
st.markdown("OECD Global Minimum Tax | FY 2024 | South Africa - Australia - Ireland")
st.markdown("---")

if page == "Dashboard":
    results = calc_globe()
    if results.empty:
        st.warning("No active entities.")
    else:
        total_income = results["GloBE_Income"].sum()
        total_tax = results["Adj_Cov_Tax"].sum()
        total_topup = results["TopUp_Tax"].sum()
        group_etr = (total_tax / total_income * 100) if total_income > 0 else 0
        k1, k2, k3, k4 = st.columns(4)
        k1.metric("Group ETR", str(round(group_etr, 2)) + "%", "Above 15%" if group_etr >= 15 else "BELOW 15%", delta_color="normal" if group_etr >= 15 else "inverse")
        k2.metric("Est. Top-up Tax", "ZAR " + str(round(total_topup, 1)) + "m", delta="Tax due" if total_topup > 0 else "Compliant", delta_color="inverse" if total_topup > 0 else "normal")
        k3.metric("GloBE Income", "ZAR " + str(round(total_income, 1)) + "m")
        k4.metric("Active Entities", len(st.session_state.entities[st.session_state.entities["Active"] == True]))
        st.markdown("---")
        st.markdown("#### Jurisdiction GloBE Summary")
        disp = results.copy()
        disp["ETR (%)"] = disp["ETR"].map(lambda x: str(round(x, 2)) + "%")
        disp["TopUp (%)"] = disp["TopUp_Rate"].map(lambda x: str(round(x, 2)) + "%")
        disp["TopUp Tax"] = disp["TopUp_Tax"].map(lambda x: "ZAR " + str(round(x, 1)) + "m")
        st.dataframe(disp[["Name", "Role", "Entities", "GloBE_Income", "Adj_Cov_Tax", "SBIE", "ETR (%)", "TopUp (%)", "TopUp Tax", "QDMTT", "Status"]], use_container_width=True, hide_index=True)
        st.markdown("---")
        st.markdown("#### Safe Harbour Assessment")
        cols = st.columns(len(results))
        for i, (_, r) in enumerate(results.iterrows()):
            with cols[i]:
                st.markdown("**" + r["Name"] + "**")
                if r["ETR"] >= 15:
                    st.success("ETR OK: " + str(round(r["ETR"], 2)) + "%")
                else:
                    st.error("ETR LOW: " + str(round(r["ETR"], 2)) + "%")
                st.info("SBIE: ZAR " + str(round(r["SBIE"], 1)) + "m excluded")
                if r["QDMTT"] == "Yes":
                    st.error("QDMTT: ZAR " + str(round(r["TopUp_Tax"], 1)) + "m due")
                else:
                    st.success("No top-up required")
        st.markdown("---")
        st.markdown("#### Entity ETR")
        ae = st.session_state.entities[st.session_state.entities["Active"] == True].copy()
        ae["ETR %"] = (ae["CoveredTaxes"] / ae["PBT"].replace(0, 1)) * 100
        st.bar_chart(ae.set_index("Entity")[["ETR %"]])

elif page == "Entities":
    st.markdown("## Entity Management")
    cfl = st.session_state.core_field_labels
    ffl = st.session_state.free_field_labels
    ffv = st.session_state.free_field_visible
    visible_free = [f"Free{i:02d}" for i in range(1, FREE_FIELD_COUNT + 1) if ffv.get(f"Free{i:02d}", False)]

    tab1, tab2 = st.tabs(["View / Edit", "Add New"])
    with tab1:
        col_cfg = {
            "Jurisdiction": st.column_config.SelectboxColumn(cfl["Jurisdiction"], options=list(JUR_INFO.keys())),
            "Type": st.column_config.SelectboxColumn(cfl["Type"], options=["Insurance", "Life Insurance", "Gen. Insurance", "Reinsurance", "Holding", "Other"]),
            "Revenue": st.column_config.NumberColumn(cfl["Revenue"], format="%.1f"),
            "PBT": st.column_config.NumberColumn(cfl["PBT"], format="%.1f"),
            "CoveredTaxes": st.column_config.NumberColumn(cfl["CoveredTaxes"], format="%.1f"),
            "DeferredTaxAdj": st.column_config.NumberColumn(cfl["DeferredTaxAdj"], format="%.1f"),
            "Payroll": st.column_config.NumberColumn(cfl["Payroll"], format="%.1f"),
            "TangibleAssets": st.column_config.NumberColumn(cfl["TangibleAssets"], format="%.1f"),
            "Active": st.column_config.CheckboxColumn("In Scope"),
        }
        for ff in visible_free:
            col_cfg[ff] = st.column_config.NumberColumn(ffl[ff], format="%.2f")
        display_cols = ["Entity", "Jurisdiction", "Type", "Revenue", "PBT", "CoveredTaxes", "DeferredTaxAdj", "Payroll", "TangibleAssets", "Active"] + visible_free
        edited = st.data_editor(
            st.session_state.entities[display_cols],
            use_container_width=True,
            num_rows="dynamic",
            column_config=col_cfg,
            hide_index=True
        )
        if st.button("Save Changes", type="primary"):
            saved = edited.reset_index(drop=True).copy()
            n_min = min(len(saved), len(st.session_state.entities))
            for i in range(1, FREE_FIELD_COUNT + 1):
                col = f"Free{i:02d}"
                if col not in saved.columns:
                    saved[col] = 0.0
                    if col in st.session_state.entities.columns:
                        for idx in range(n_min):
                            saved.at[idx, col] = st.session_state.entities.at[idx, col]
            st.session_state.entities = saved
            st.success("Saved!")
            st.rerun()
        if visible_free:
            st.info(f"{len(visible_free)} custom field(s) shown. Manage fields in **Settings**.")
        else:
            st.info("No custom fields enabled. Go to **Settings** to enable and label your 20 free fields.")

    with tab2:
        c1, c2, c3 = st.columns(3)
        with c1:
            n_name = st.text_input(cfl["Entity"])
            n_jur = st.selectbox(cfl["Jurisdiction"], list(JUR_INFO.keys()))
            n_type = st.selectbox(cfl["Type"], ["Insurance", "Life Insurance", "Gen. Insurance", "Reinsurance", "Holding", "Other"])
        with c2:
            n_rev = st.number_input(cfl["Revenue"], value=0.0)
            n_pbt = st.number_input(cfl["PBT"], value=0.0)
            n_tax = st.number_input(cfl["CoveredTaxes"], value=0.0)
        with c3:
            n_dtx = st.number_input(cfl["DeferredTaxAdj"], value=0.0)
            n_pay = st.number_input(cfl["Payroll"], value=0.0)
            n_tan = st.number_input(cfl["TangibleAssets"], value=0.0)
        free_vals = {}
        if visible_free:
            st.markdown("**Custom Fields**")
            ff_cols = st.columns(min(4, len(visible_free)))
            for k, ff in enumerate(visible_free):
                with ff_cols[k % len(ff_cols)]:
                    free_vals[ff] = st.number_input(ffl[ff], value=0.0, key=f"add_{ff}")
        if st.button("Add Entity", type="primary"):
            if n_name:
                free_defaults = {f"Free{i:02d}": 0.0 for i in range(1, FREE_FIELD_COUNT + 1)}
                free_defaults.update(free_vals)
                new_row = pd.DataFrame([{**{"Entity": n_name, "Jurisdiction": n_jur, "Type": n_type, "Revenue": n_rev, "PBT": n_pbt, "CoveredTaxes": n_tax, "DeferredTaxAdj": n_dtx, "Payroll": n_pay, "TangibleAssets": n_tan, "Active": True}, **free_defaults}])
                st.session_state.entities = pd.concat([st.session_state.entities, new_row], ignore_index=True)
                st.success("Added " + n_name)
                st.rerun()

elif page == "Transactions":
    st.markdown("## Intercompany Transaction Register")
    tab1, tab2 = st.tabs(["Register", "Add New"])
    entity_names = list(st.session_state.entities["Entity"])
    with tab1:
        edited_tx = st.data_editor(
            st.session_state.transactions,
            use_container_width=True,
            num_rows="dynamic",
            column_config={
                "From": st.column_config.SelectboxColumn("From", options=entity_names),
                "To": st.column_config.SelectboxColumn("To", options=entity_names),
                "Type": st.column_config.SelectboxColumn("Type", options=["Reinsurance Premium", "Ceding Commission", "Management Fee", "Royalty / IP", "Loan Interest", "Capital Contribution", "Dividend", "Claims Reimbursement"]),
                "TP_Method": st.column_config.SelectboxColumn("TP Method", options=["CUP", "TNMM", "CPS", "PSM", "CPM"]),
                "Amount": st.column_config.NumberColumn("Amount (ZARm)", format="%.1f"),
                "ArmsLength": st.column_config.CheckboxColumn("Arms Length"),
            },
            hide_index=True
        )
        if st.button("Save Transactions", type="primary"):
            st.session_state.transactions = edited_tx
            st.success("Saved!")
        flagged = st.session_state.transactions[st.session_state.transactions["ArmsLength"] == False]
        if len(flagged) > 0:
            st.warning(str(len(flagged)) + " transaction(s) flagged for review.")
    with tab2:
        c1, c2 = st.columns(2)
        with c1:
            tx_desc = st.text_input("Description")
            tx_from = st.selectbox("From", entity_names)
            tx_to = st.selectbox("To", entity_names, index=min(3, max(0, len(entity_names) - 1)))
            tx_amt = st.number_input("Amount (ZARm)", value=0.0)
        with c2:
            tx_type = st.selectbox("Type", ["Reinsurance Premium", "Ceding Commission", "Management Fee", "Royalty / IP", "Loan Interest", "Capital Contribution", "Dividend", "Claims Reimbursement"])
            tx_tp = st.selectbox("TP Method", ["CUP", "TNMM", "CPS", "PSM", "CPM"])
            tx_arm = st.checkbox("Arms Length", value=True)
        if st.button("Add Transaction", type="primary"):
            new_tx = pd.DataFrame([{"Description": tx_desc, "From": tx_from, "To": tx_to, "Amount": tx_amt, "Type": tx_type, "TP_Method": tx_tp, "ArmsLength": tx_arm}])
            st.session_state.transactions = pd.concat([st.session_state.transactions, new_tx], ignore_index=True)
            st.success("Added!")
            st.rerun()

elif page == "Upload Data":
    st.markdown("## Upload and Auto-Populate")
    st.markdown("Upload trial balance or statutory accounts CSV/Excel to populate entity fields.")
    template_df = pd.DataFrame({"Entity": ["Example Co"], "Jurisdiction": ["ZA"], "Type": ["Insurance"], "Revenue": [0.0], "PBT": [0.0], "CoveredTaxes": [0.0], "DeferredTaxAdj": [0.0], "Payroll": [0.0], "TangibleAssets": [0.0]})
    st.download_button("Download CSV Template", data=template_df.to_csv(index=False), file_name="pillar2_template.csv", mime="text/csv")
    st.markdown("---")
    uploaded = st.file_uploader("Upload CSV or Excel", type=["csv", "xlsx", "xls"])
    if uploaded:
        try:
            df_raw = pd.read_csv(uploaded) if uploaded.name.endswith(".csv") else pd.read_excel(uploaded)
            st.dataframe(df_raw.head(20), use_container_width=True)
            targets = {
                "Entity": ["entity", "name", "company"],
                "Jurisdiction": ["jurisdiction", "jur", "country"],
                "Type": ["type", "entity type"],
                "Revenue": ["revenue", "turnover", "gross premium"],
                "PBT": ["pbt", "profit before tax", "profit"],
                "CoveredTaxes": ["covered taxes", "tax", "income tax"],
                "DeferredTaxAdj": ["deferred", "deferred tax"],
                "Payroll": ["payroll", "staff costs", "wages"],
                "TangibleAssets": ["tangible", "ppe", "fixed assets"],
            }
            col_map = {}
            for target, aliases in targets.items():
                for col in df_raw.columns:
                    if any(a in col.lower() for a in aliases):
                        col_map[target] = col
                        break
            final_map = {}
            for target in targets:
                detected = col_map.get(target, "")
                options = ["(skip)"] + list(df_raw.columns)
                idx = (list(df_raw.columns).index(detected) + 1) if detected in df_raw.columns else 0
                chosen = st.selectbox(target, options, index=idx, key="map_" + target)
                if chosen != "(skip)":
                    final_map[target] = chosen
            if st.button("Import Data", type="primary"):
                imported = pd.DataFrame()
                for target, src in final_map.items():
                    imported[target] = df_raw[src]
                for col, default in [("Active", True), ("Type", "Insurance"), ("DeferredTaxAdj", 0.0), ("Payroll", 0.0), ("TangibleAssets", 0.0)]:
                    if col not in imported.columns:
                        imported[col] = default
                for nc in ["Revenue", "PBT", "CoveredTaxes", "DeferredTaxAdj", "Payroll", "TangibleAssets"]:
                    if nc in imported.columns:
                        imported[nc] = pd.to_numeric(imported[nc], errors="coerce").fillna(0.0)
                for i in range(1, FREE_FIELD_COUNT + 1):
                    imported[f"Free{i:02d}"] = 0.0
                st.session_state.entities = pd.concat([st.session_state.entities, imported], ignore_index=True)
                st.success("Imported " + str(len(imported)) + " entities!")
                st.rerun()
        except Exception as e:
            st.error("Error: " + str(e))

elif page == "AI Templates":
    st.markdown("## AI Template Generator")
    TEMPLATES = ["GloBE Information Return (GIR)", "QDMTT Calculation Worksheet", "IIR Allocation Memorandum", "Transfer Pricing Impact Analysis", "Insurance-Specific Deferred Tax Adjustment", "SBIE Workpaper", "Transitional Safe Harbour Assessment", "CbCR Reconciliation", "Board Tax Governance Report", "BEPS Action 13 Master File Summary"]
    selected = st.selectbox("Select Template", TEMPLATES)
    generate = st.button("Generate Template", type="primary")
    st.markdown("Quick select:")
    cols = st.columns(3)
    for i, t in enumerate(TEMPLATES):
        if cols[i % 3].button(t, key="tpl_" + str(i), use_container_width=True):
            selected = t
            generate = True
    if generate:
        summary = get_summary()
        entity_list = ", ".join([r["Entity"] + " (" + r["Jurisdiction"] + ")" for _, r in st.session_state.entities[st.session_state.entities["Active"] == True].iterrows()])
        tx_list = ", ".join([r["Description"] + " ZAR" + str(r["Amount"]) + "m" for _, r in st.session_state.transactions.iterrows()])
        prompt = build_template_prompt(selected, entity_list, summary, tx_list)
        with st.spinner("Generating " + selected + "..."):
            output = call_claude(prompt)
        st.markdown("---")
        st.markdown("### " + selected)
        st.text_area("Output", value=output, height=500, label_visibility="collapsed")
        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button("Download as TXT", data=output, file_name=selected.replace(" ", "_") + ".txt", mime="text/plain")
        with dl2:
            word_bytes = generate_word_doc(selected, output)
            st.download_button("Download as Word (.docx)", data=word_bytes, file_name=selected.replace(" ", "_") + ".docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif page == "Benchmarking":
    st.markdown("## Global Best Practice Benchmarking")
    PRESETS = [
        "How should OUTsurance Ireland handle the QDMTT vs IIR interaction for FY2024?",
        "What are the insurance-specific deferred tax adjustments under GloBE Article 4.4?",
        "Benchmark OUTsurance ETR against global insurance peers such as Zurich and Munich Re",
        "What transitional safe harbour elections should OUTsurance make for 2024 to 2026?",
        "How do intragroup reinsurance flows between ZA and IE affect GloBE income allocation?",
        "What documentation is required for the GloBE Information Return in South Africa?",
        "Explain the UTPR rules and whether OUTsurance is exposed as South African HQ",
        "What are the Pillar II implications of OUTsurance deferred acquisition costs?",
    ]
    st.markdown("Suggested questions:")
    cols = st.columns(2)
    selected_preset = None
    for i, p in enumerate(PRESETS):
        if cols[i % 2].button(p, key="preset_" + str(i), use_container_width=True):
            selected_preset = p
    st.markdown("---")
    query = st.text_area("Or type your question:", value=selected_preset or "", height=80)
    if st.button("Get AI Analysis", type="primary") and query:
        summary = get_summary()
        prompt = build_benchmarking_prompt(query, summary)
        with st.spinner("Analysing..."):
            answer = call_claude(prompt)
        st.markdown("---")
        st.markdown("### AI Analysis")
        st.markdown(answer)
        dl1, dl2 = st.columns(2)
        with dl1:
            st.download_button("Save Analysis as TXT", data=answer, file_name="pillar2_analysis.txt", mime="text/plain")
        with dl2:
            word_bytes = generate_word_doc("Pillar II GloBE Analysis", answer)
            st.download_button("Save Analysis as Word (.docx)", data=word_bytes, file_name="pillar2_analysis.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.markdown("---")
    st.markdown("### Jurisdiction Reference")
    c1, c2, c3 = st.columns(3)
    with c1:
        with st.expander("South Africa", expanded=True):
            st.markdown("- Legislation: TLAB 2023\n- QDMTT: 1 Jan 2024\n- STR: 27%\n- Role: UPE / IIR filer\n- Filing: 12 months post year-end")
    with c2:
        with st.expander("Australia", expanded=True):
            st.markdown("- Legislation: Treasury Laws Amendment Act 2024\n- STR: 30%\n- No top-up expected\n- Safe harbour: 2024-2026")
    with c3:
        with st.expander("Ireland", expanded=True):
            st.markdown("- Legislation: Finance Act 2024\n- STR: 12.5%\n- QDMTT top-up to 15% required\n- PRIMARY RISK for OUTsurance")

elif page == "Settings":
    st.markdown("## Settings & Field Configuration")
    tab1, tab2 = st.tabs(["Core Fields", "Free Fields (20)"])

    with tab1:
        st.markdown("### Core GloBE Field Labels")
        st.markdown("Customise the display labels for the core GloBE entity fields. These fields are required for tax calculations and cannot be disabled.")
        cfl = st.session_state.core_field_labels
        c1, c2, c3 = st.columns(3)
        fields = list(cfl.keys())
        for j, field in enumerate(fields):
            col = [c1, c2, c3][j % 3]
            with col:
                new_lbl = st.text_input(f"Label for '{field}'", value=cfl[field], key=f"core_{field}")
                st.session_state.core_field_labels[field] = new_lbl
        if st.button("Save Core Labels", type="primary", key="save_core"):
            st.success("Core field labels updated!")

    with tab2:
        st.markdown("### Free Quantitative Fields")
        st.markdown(
            "Enable and name up to **20 additional quantitative fields** per company. "
            "Use these to capture any additional financial metrics beyond the standard GloBE fields — "
            "e.g. *Technical Provisions*, *DAC Balance*, *Equalisation Reserve*, *Gross Written Premium*, etc."
        )
        st.markdown("---")
        ffl = st.session_state.free_field_labels
        ffv = st.session_state.free_field_visible
        enabled_count = sum(1 for v in ffv.values() if v)
        st.info(f"{enabled_count} of {FREE_FIELD_COUNT} free fields currently enabled.")
        cols = st.columns(4)
        for i in range(1, FREE_FIELD_COUNT + 1):
            field = f"Free{i:02d}"
            col = cols[(i - 1) % 4]
            with col:
                st.markdown(f"**Field {i}**")
                vis = st.checkbox("Enable", value=ffv[field], key=f"vis_{field}")
                lbl = st.text_input("Label", value=ffl[field], key=f"lbl_{field}", disabled=not vis)
                st.session_state.free_field_visible[field] = vis
                if vis:
                    st.session_state.free_field_labels[field] = lbl
                st.markdown("")
        if st.button("Save Free Field Settings", type="primary", key="save_free"):
            st.success("Free field settings saved! Go to **Entities** to see and edit the enabled fields.")
